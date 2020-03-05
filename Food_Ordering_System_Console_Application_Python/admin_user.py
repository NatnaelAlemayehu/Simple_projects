# Author: Natnael Alemayehu / Mahalinoro Razafimanjato

# Importing the class files needed
import docx

from user_class import User
import food_class
import drink_class
import sys
import main_application
import item_class

# Dictionary that will contain the admin users after creation
admin_user = {}



# This is the class admin user with the instance variables and the __init__ method


class AdminUser(User):
    def __init__(self, name):
        super().__init__(name)


def view_items():
    # This function will display the food items and the drinks item to the user
    print("You have this foods: ")
    for food in food_class.food_item.keys():
        print(food)
    print("You have this drinks: ")
    for drink in drink_class.drink_item.keys():
        print(drink)
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin viewed food and drink items')
    doc.save("admin_user_session.docx")
    admin_menu()



# This function will add new food item in the storage
def add_food_item(food_name, food_price, food_size, food_description):
    if food_name == "" or food_price == "" or food_size == "" or food_description == "":
        print("input can not be empty")
        admin_menu()
    food_class.food_item.update({food_name: food_class.Food(food_name,
                                                            food_price, food_size, food_description)})
    print("Successfully added food item")
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin added {} to food list'.format(food_name))
    doc.save("admin_user_session.docx")
    admin_menu()

# This function will add new drink in the storage


def add_drink_item(drink_name, drink_price, drink_volume, drink_stock, drink_description):
    if drink_name == "" or drink_price == "" or drink_volume == "" or drink_stock == "" or drink_description == "":
        print("input can not be empty")
        admin_menu()
    drink_class.drink_item.update({drink_name: drink_class.Drink(
        drink_name, drink_price, drink_volume, drink_stock, drink_description)})
    print("Successfully added food item")
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin added {} to drink list'.format(drink_name))
    doc.save("admin_user_session.docx")
    admin_menu()


# This function will remove a selected food item
def remove_food_item(food_to_delete):
    if food_to_delete == "":
        print("input can not be empty")
        admin_menu()
    if food_to_delete in food_class.food_item.keys():
        food_class.food_item.pop(food_to_delete)
        print("Item removed successfully")
        doc = docx.Document('admin_user_session.docx')
        doc.add_paragraph('Admin removed {} from food list'.format(food_to_delete))
        doc.save("admin_user_session.docx")
        admin_menu()
    else:
        print("food not found")
        admin_menu()

# This function will remove a selected drink item
def remove_drink_item(drink_to_delete):
    if drink_to_delete == "":
        print("input can not be empty")
        admin_menu()
    if drink_to_delete in drink_class.drink_item.keys():
        drink_class.drink_item.pop(drink_to_delete)
        print("Item removed successfully")
        doc = docx.Document('admin_user_session.docx')
        doc.add_paragraph(
            'Admin removed {} from drink list'.format(drink_to_delete))
        doc.save("admin_user_session.docx")
        admin_menu()
    else:
        print("drink not found")
        admin_menu()



# This function will update the information of a selected food item
def update_food_information(food_to_update, update_choice):
    if food_to_update == "" or update_choice == "":
        print("input can not be empty")
        admin_menu()
    if update_choice == "name":
        updated_name = input(
            "What do you want to rename it as? \n")
        try:
            updated_name = int(updated_name)
            print("Food name can't be an integer")
            admin_menu()
        except:
            food_class.food_item[food_to_update].name = updated_name
            food_class.food_item[updated_name] = food_class.food_item[food_to_update]
    elif update_choice == "price":
        updated_price = input("What do you want to update it as? \n")
        try:
            updated_price = int(updated_price)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        food_class.food_item[food_to_update].price = updated_price
    elif update_choice == "size":
        updated_size = input("What do you want to update it as? \n")
        try:
            updated_size = int(updated_size)
            print("Size can't be an integer")
            admin_menu()
        except:
            food_class.food_item[food_to_update].size = updated_size
    elif update_choice == "description":
        updated_description = input(
            "What do you want to update it as? \n")
        try:
            updated_description = int(updated_description)
            print("Description can't be an integer")
            admin_menu()
        except:
            food_class.food_item[food_to_update].description = updated_description
    else:
        print("You can only put name, price, size or description")
        admin_menu()
    print("Item updated successfully")
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin updated {} to {} from food list'.format(
        food_to_update, update_choice))
    doc.save("admin_user_session.docx")
    admin_menu()



# This function will update the information of a selected drink item
def update_drink_information(drink_to_update, update_choice):
    if drink_to_update == "" or update_choice == "":
        print("input can not be empty")
        admin_menu()
    if update_choice == "name":
        updated_name = input(
            "What do you want to rename it as? \n")
        try:
            updated_name = int(updated_name)
            print("Food name can't be an integer")
            admin_menu()
        except:
            drink_class.drink_item[drink_to_update].name = updated_name
            drink_class.drink_item[updated_name] = drink_class.drink_item.pop(drink_to_update)
    elif update_choice == "price":
        updated_price = input("What do you want to update it as? \n")
        try:
            updated_price = int(updated_price)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        drink_class.drink_item[drink_to_update].price = updated_price
    elif update_choice == "volume":
        updated_volume = input("What do you want to update it as? \n")
        try:
            updated_volume = int(updated_volume)
        except ValueError:
            print("Only strings allowed")
            admin_menu()
        drink_class.drink_item[drink_to_update].volume = updated_volume
    elif update_choice == "description":
        updated_description = input(
            "What do you want to update it as? \n")
        drink_class.drink_item[drink_to_update].description = updated_description
    print("Item updated successfully")
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin updated {} to {} from food list'.format(
        drink_to_update, update_choice))
    doc.save("admin_user_session.docx")
    admin_menu()


# This function will generate the revenue made from drink and food sold so far
def revenue_generated():
    food_revenue = 0
    for food in food_class.food_item.keys():
        food_revenue += food_class.food_item[food].price * food_class.food_item[food].item_count_delivered
    print("Total revenue from food is", food_revenue)
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin checked revenue from foods which was equal to {}'.format(
        food_revenue))
    doc.save("admin_user_session.docx")
    drink_revenue = 0
    for drink in drink_class.drink_item.keys():
        drink_revenue += drink_class.drink_item[drink].price * drink_class.drink_item[drink].item_count_delivered
    print("Total revenue from drink is", drink_revenue)
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin checked revenue from drinks which was equal to {}'.format(
        drink_revenue))
    doc.save("admin_user_session.docx")
    net_revenue = 0
    net_revenue = food_revenue + drink_revenue
    print("Total revenue is", net_revenue)
    doc = docx.Document('admin_user_session.docx')
    doc.add_paragraph('Admin checked total revenue from drinks which was equal to {}'.format(
        net_revenue))
    doc.save("admin_user_session.docx")
    admin_menu()




# This function will display the number of order for food and drinks item
def order_number():
    food_item_count = 0
    drink_item_count = 0
    doc = docx.Document('admin_user_session.docx')
    for food in food_class.food_item.keys():
        food_item_count += food_class.food_item[food].item_count_delivered
    for drink in drink_class.drink_item.keys():
         drink_item_count += drink_class.drink_item[drink].item_count_delivered
    print("{} orders for both food and drink".format(
            str(food_item_count + drink_item_count)))

    doc.add_paragraph('Admin checked order number for both food and drink which is equal to {}'.format(
            str(food_item_count + drink_item_count)))
    doc.save("admin_user_session.docx")
    admin_menu()


# this functions takes in users document file which contains names of places and 
# appends them to the fast_delivery_place_list
        
def update_fast_delivery_places(file_name):
    doc = docx.Document('admin_user_session.docx')
    if file_name != "":
        with open(file_name, "r") as f:
            for line in f:
                item_class.fast_time_delivery_places.extend(line.split())
        print("Successfully added items")
        doc.add_paragraph("user added place names to fast delivery list")
        doc.save("admin_user_session.docx")        
        admin_menu()
    else:
        print("empty name passed")
        admin_menu()


# this functions takes in users document file which contains names of places and
# appends them to the medium_delivery_place_list

def update_medium_delivery_places(file_name):
    doc = docx.Document('admin_user_session.docx')
    if file_name != "":
        with open(file_name, "r") as f:
            for line in f:
                item_class.medium_time_delivery_places.extend(line.split())
        print("Successfully added items")
        doc.add_paragraph("user added place names to fast delivery list")
        doc.save("admin_user_session.docx")
        admin_menu()
    else:
        print("empty name passed")
        admin_menu()
    


# This function is the main menu for admin user
# The admin menu will let the user view item, add new item, remove item, update information, see revenue and see
# order numbers as well as they can log out of the system if they wish

def admin_menu():
    doc = docx.Document('admin_user_session.docx')
    print('----------- Welcome to E-MEAL MENU ---------')
    print("Press 1 to view your item in the storage ")
    print("Press 2 to add new item in the storage ")
    print("Press 3 to remove item in the storage")
    print("Press 4 to update information about a specific item")
    print("Press 5 to see revenue generated")
    print("Press 6 to see order numbers for items")
    print("Press U to import delivery place names from doc/docx file")
    print("Press L to logout of your account")
    print("Press q or Q to quit.")
    user_choice = input(
        "Type in the corresponding number for you want to do \n")
    if user_choice == "1":
        doc.add_paragraph('User pressed 1 to view items in the storage.')
        doc.save("admin_user_session.docx")
        view_items()
    elif user_choice == "2":
        doc.add_paragraph(
            'User pressed 2 to add new item in the storage.')
        doc.save("admin_user_session.docx")
        print("What do you want to add?")
        user_input = input("Press 1 to add food and 2 to add drink \n")
        if user_input == "1":
            doc.add_paragraph('User pressed 1 to add new food.')
            food_name = input("Food name \n")
            doc.add_paragraph(
                'User inserted {} as food name.'.format(food_name))
            food_price = input("price of food \n")
            doc.add_paragraph(
                'User inserted {} as food price.'.format(food_price))
            doc.save("admin_user_session.docx")
            try:
                food_price = int(food_price)
            except ValueError:
                doc.add_paragraph(
                    'User got errors for inserting not integer price for food.')
                print("Integer only")
                doc.save("admin_user_session.docx")
                admin_menu()
            food_size = input("Size \n")
            doc.add_paragraph(
                'User inserted {} as food size.'.format(food_size))
            food_description = input("description \n")
            doc.add_paragraph(
                'User inserted {} as food description.'.format(food_description))
            doc.save("admin_user_session.docx")
            add_food_item(food_name, food_price, food_size, food_description)
        elif user_input == "2":
            doc.add_paragraph('User pressed 2 to add new drink.')
            drink_name = input("Drink name \n")
            doc.add_paragraph(
                'User inserted {} as drink name.'.format(drink_name))
            drink_price = input("price of drink \n")
            doc.add_paragraph('User inserted {} as drink price.'.format(drink_price))
            drink_volume = input("Drink volume \n")
            doc.add_paragraph(
                'User inserted {} as drink volume.'.format(drink_volume))
            drink_stock = input("Stock \n")
            doc.add_paragraph(
                'User inserted {} as drink stock.'.format(drink_stock))
            doc.save("admin_user_session.docx")
            try:
                drink_price = int(drink_price)
                drink_volume = int(drink_volume)
                drink_stock = int(drink_stock)
            except ValueError:
                doc.add_paragraph(
                    'User got errors for inserting non integer value for price, volume, stock.')
                print("integer only")
                doc.save("admin_user_session.docx")
                admin_menu()
            drink_description = input("Drink description \n")
            doc.add_paragraph(
                'User inserted {} as drink description.'.format(drink_description))
            doc.save("admin_user_session.docx")
            add_drink_item(drink_name, drink_price, drink_volume,
                           drink_stock, drink_description)
            doc.add_paragraph(
                'User added {} successfully in the Drink Item.'.format(drink_name))
            doc.save("admin_user_session.docx")
        else:
            print("Error try again")
            doc.add_paragraph(
                'User inserted wrong input')
            doc.save("admin_user_session.docx")
            admin_menu()
    elif user_choice == "3":
        doc.add_paragraph(
            'User pressed {} to remove item.'.format(user_choice))
        print("What do you want to delete")
        user_input = input(
            "press 1 to delete food items or 2 for drink items \n")
        if user_input == "1":
            doc.add_paragraph(
                'User pressed {} to delete food items.'.format(user_input))
            doc.save("admin_user_session.docx")
            for item in food_class.food_item.keys():
                print(item)
            food_to_delete = input("type in name of food to delete \n")
            doc.add_paragraph(
                'User inserted {} to delete.'.format(food_to_delete))
            remove_food_item(food_to_delete)
            doc.add_paragraph(
                'User deleted {} definitely.'.format(food_to_delete))
            doc.save("admin_user_session.docx")
        elif user_input == "2":
            doc.add_paragraph(
                'User pressed {} to delete drink items.'.format(user_input))
            for item in drink_class.drink_item.keys():
                print(item)
            drink_to_delete = input("type in name of drink to delete \n")
            doc.add_paragraph(
                'User inserted {} to delete drink items.'.format(drink_to_delete))
            remove_drink_item(drink_to_delete)
            doc.add_paragraph(
                'User deleted {} definitely.'.format(drink_to_delete))
            doc.save("admin_user_session.docx")
        else:
            print("Error try again")
            doc.add_paragraph(
                'User received error for wrong input.')
            doc.save("admin_user_session.docx")
            admin_menu()
    elif user_choice == "4":
        doc.add_paragraph(
            'User pressed {} to update items information.'.format(user_choice))
        print("What do you want to update")
        user_input = input(
            "press 1 to update food items or 2 to update drink items \n")
        if user_input == "1":
            doc.add_paragraph(
                'User pressed {} to update food items.'.format(user_input))
            doc.save("admin_user_session.docx")
            for item in food_class.food_item.keys():
                print(item)
            food_to_update = input("type in name of food to update \n")
            doc.add_paragraph(
                'User selected {} to update.'.format(food_to_update))
            doc.save("admin_user_session.docx")
            if food_to_update in food_class.food_item.keys():
                update_choice = input(
                    "What do you want to update? name, price, size? \n")
                update_food_information(food_to_update, update_choice)
        elif user_input == "2":
            doc.add_paragraph(
                'User pressed {} to update drink items.'.format(user_input))
            doc.save("admin_user_session.docx")
            for item in drink_class.drink_item.keys():
                print(item)
            drink_to_update = input("type in name of food to delete \n")
            doc.add_paragraph(
                'User selected {} to update.'.format(drink_to_update))
            doc.save("admin_user_session.docx")
            if drink_to_update in drink_class.drink_item.keys():
                update_choice = input(
                    "What do you want to update? \n name, price, volume? \n")
                update_drink_information(drink_to_update, update_choice)
        else:
            print("Error try again")
            doc.add_paragraph(
                'User gets error for wrong input.')
            doc.save("admin_user_session.docx")
        admin_menu()
    elif user_choice == "5":
        doc.add_paragraph(
            'User pressed {} to check revenue.'.format(user_choice))
        doc.save("admin_user_session.docx")
        revenue_generated()
    elif user_choice == "6":
        doc.add_paragraph(
            'User pressed {} to see order number.'.format(user_choice))
        doc.save("admin_user_session.docx")
        order_number()
    elif user_choice == "U":
        user_input = input("Press 1 to add places to place list \n ")
        try:
            user_input = int(user_input)  
            second_choice = input('''Press 1 to update fast delivery places
Press 2 to update medium delivery places \n''') 
            try:
                second_choice = int(second_choice)
                if second_choice == 1:
                    file_name = input("What is the name of your doc file (include extension) \n")
                    update_fast_delivery_places(file_name)
                elif second_choice == 2:
                    file_name = input(
                        "What is the name of your doc file (include extension) \n")
                    update_medium_delivery_places(file_name)
            except:
                print("integer only")  
                admin_menu()       
        except:
            print("integer only")
            admin_menu()
        item_class.fast_time_delivery_places
    elif user_choice == "L":
        doc.add_paragraph('User pressed {} to Log out.'.format(user_choice))
        doc.save("admin_user_session.docx")        
        main_application.menu()
    elif user_choice == "q" or "Q":
        doc.add_paragraph(
            'User pressed {} quit the program.'.format(user_choice))
        doc.save("admin_user_session.docx")
        sys.exit()
    else:
        print("Error, try again!")
        doc.add_paragraph('User pressed wrong option.')
        doc.save("admin_user_session.docx")
        admin_menu()


if __name__ == '__main__':
    admin_menu()
