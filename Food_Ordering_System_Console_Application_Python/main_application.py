# Author: Natnael Alemayehu / Mahalinoro Razafimanjato

# Importing the class needed
import regular_user_class
import docx
import admin_user
import datetime

# Main program workflow
# This function contains the menu and the user input to choose and authenticate as admin or regular
# If the user doesn't have an account, it will be redirected to the creation of account
# If the user already gave an account, it will be redirected to the main menu corresponding to the type of user


doc = docx.Document()


def menu():
    global doc
    print(" | Welcome to the E-Meal |")
    doc.add_heading('Admin User Session', 1)
    today = datetime.datetime.today()
    doc.add_heading('Date: {}'.format(today), 2)
    user_type = input('''Press 1 to login or sign up as an Admin
Press 2 to login or sign up as customer \n''')
    if user_type == "1":
        doc.add_paragraph('User pressed 1 to login or sign up as an Admin.')
        user_status = input('''Press 1 to Signup
Press 2 to Login \n''')
        if user_status == "1":
            doc.add_paragraph('User pressed 1 to Signup.')
            create_admin_account()
        elif user_status == "2":
            doc.add_paragraph('User pressed 2 to Login.')
            user_name = input("What is your username \n")
            doc.add_paragraph('User insert {} as username to login'.format(user_name))
            login(user_name)
        else:
            print('Please press 1 or 2! Try again:')
            menu()
    elif user_type == "2":
        user_status = input('''Press 1 to Signup
Press 2 to Login \n''')
        doc.add_paragraph('User pressed 2 to login or sign up as Regular User.')
        if user_status == "1":
            doc.add_paragraph('User pressed 1 to Signup.')
            create_regular_account()
        elif user_status == "2":
            doc.add_paragraph('User pressed 2 to Login.')
            user_name = input("What is your username \n")
            doc.add_paragraph('User inserted {} as username to login.'.format(user_name))
            login(user_name)
        else:
            print('Please press 1 or 2! Try again: ')
            menu()
    else:
        print('Please press 1 or 2 only! Try again:')
        doc.add_paragraph('User pressed {} and get an error message.'.format(user_type))
        menu()
    doc.save('admin_user_session.docx')


# This function creates a new regular account and it will be appended in the dictionary
def create_regular_account():
    global doc
    username = input("What is your name? \n").lower()
    location = input("Where is your location \n").lower()
    wallet = input("How much do you want to put in your wallet? \n")
    try:
        wallet = int(wallet)
    except ValueError:
        print('Only integer allowed for wallet.')
        create_regular_account()
    regular_user_class.regular_user.update({username: regular_user_class.RegularUser(username, location, wallet)})
    print("Account created successfully")
    print("Your user name is {}".format(username))
    doc.add_paragraph('User is creating a new account:')
    doc.add_paragraph('User inserted {} as username.'.format(username))
    doc.add_paragraph('User inserted {} as location.'.format(location))
    doc.add_paragraph('User inserted {} as amount in the wallet.'.format(wallet))
    doc.add_paragraph('User {} created and added into Regular User.'.format(username))
    doc.save('regular_user_session.docx')
    login(username)


# This function will verify it the user is admin ore regular
# If it is a regular user, then it will go directly into the ordering menu
# If it is an admin user, then it will go directly into the admin menu
def login(username):
    global doc
    if username in regular_user_class.regular_user.keys():
        doc.add_paragraph('User {} redirected to Regular User Menu.'.format(username))
        doc.save('admin_user_session.docx')
        regular_user_class.make_order()
        regular_user_class.make_payments(username)
        regular_user_class.user_review()
    elif username in admin_user.admin_user.keys():
        doc.add_paragraph('User {} redirected to Admin User Menu.'.format(username))
        doc.save('admin_user_session.docx')
        admin_user.admin_menu()
    else:
        print("User not found. Create an account first")
        doc.add_paragraph('Username {} is not found in the database.'.format(username))
        doc.save('admin_user_session.docx')
        menu()

# This function will create an admin account
def create_admin_account():
    global doc
    username = input("Insert name: ").lower()
    doc.add_paragraph('User is creating a new account:')
    if username == "":
        print('Empty name not allowed.')
        doc.add_paragraph('User insert empty input for name and get an error.')
        create_admin_account()
    else:
        admin_user.admin_user.update({username: admin_user.AdminUser(username)})
        print("Account created successfully")
        print("Your user name is {}".format(username))
        doc.add_paragraph('User inserted {} as username.'.format(username))
        doc.add_paragraph('User {} created and added to Admin User.'.format(username))
        doc.add_paragraph('User {} redirected to Admin User Menu.'.format(username))
        doc.save('admin_user_session.docx')
        admin_user.admin_menu()


if __name__ == "__main__":
    menu()


