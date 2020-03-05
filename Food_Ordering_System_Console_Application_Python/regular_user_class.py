# Author: Mahalinoro Razafimanjato
# Import the user class parent, drink and food for functions
from user_class import User
import drink_class
import food_class
import item_class
import docx
import datetime
import main_application
# Dictionary list for the regular user
regular_user = {}

# Global variables
user_choice_item = ""
amount = ""
itemtype = ""
item_count_delivered = 0

doc = docx.Document()


# Regular User class which is the child of the User class
class RegularUser(User):
    # Method __init__ for the regular user
    def __init__(self, name, location, wallet):
        super().__init__(name)
        self.location = location
        self.wallet = wallet

    def __str__(self):
        pass

    # Method to update information for the instance variables
    def update_information(self, new_name, new_location):
        self.name = new_name
        self.location = new_location
        print("Information updated successfully.")
        return self.name, self.location

    # Method to add money to the wallet
    def add_money(self, updated_amount):
        self.wallet = self.wallet + updated_amount
        return self.wallet

    # Method to subtract money from the wallet
    def subtract_money(self, total_price):
        self.wallet = self.wallet - total_price
        return self.wallet


# Functions Section
def make_order():
    # This function will help the user to order a specific item
    # If the customer wants food/drink
    global doc
    doc.add_heading('Regular User Session', 1)
    today = datetime.datetime.today()
    doc.add_heading('Date: {}'.format(today), 2)
    doc.save('regular_user_session.docx')
    global user_choice_item
    global amount
    global itemtype
    itemtype = input(
        "Which type of item do you wish to order [drink/food]:  \n").lower()
    if itemtype == "food":
        doc.add_paragraph('User selected food as item to order.')
        print("-------------- FOOD MENU ----------------")
        food_class.view_food_items()
        user_choice_item = input("Type in your food order:  \n")
        if user_choice_item == "" or type(user_choice_item) == int: 
            print("Invalid input")
            doc.add_paragraph('User inserted invalid input')
            doc.save('regular_user_session.docx')            
            make_order()
        else:
            doc.add_paragraph('User inserted {} as food to order.'.format(user_choice_item))
            doc.save('regular_user_session.docx')
            amount = input("Insert the quantity [ex: 1,2] :  \n")
            try:
                amount = int(amount)
            except ValueError:
                print('Only integer allowed. Try again!')
                make_order()
            doc.add_paragraph('User inserted {} as quantity of food to order.'.format(amount))
            doc.save('regular_user_session.docx')
            confirmation = input("Are you confirming your order? [yes/no]:  \n").lower()
            if confirmation == "yes":
                print("Order confirmed!")
                doc.add_paragraph('User inserted {} to confirm the order.'.format(confirmation))
                doc.save('regular_user_session.docx')
            else:
                doc.add_paragraph('User inserted {} to cancel order.'.format(confirmation))
                doc.save('regular_user_session.docx')
                make_order()
    elif itemtype == "drink":
        doc.add_paragraph('User selected drink as item to order.')
        print("--------------- DRINK MENU ---------------")
        drink_class.view_drink_items()
        user_choice_item = input("Your drink order:  \n")
        if user_choice_item == "" or type(user_choice_item) == int:
            print("Invalid input")
            doc.add_paragraph('User inserted invalid input')
            doc.save('regular_user_session.docx')            
            make_order()
        else:
            doc.add_paragraph('User inserted {} as drink to order.'.format(user_choice_item))
            doc.save('regular_user_session.docx')
            amount = input("Insert the quantity [ex: 1,2] :  \n")
            try:
                amount = int(amount)
            except ValueError:
                print('Only integer allowed. Try again!')
                make_order()
            doc.add_paragraph('User inserted {} as quantity drink to order.'.format(amount))
            doc.save('regular_user_session.docx')
            confirmation = input("Are you confirming your order? [yes/no]:  \n").lower()
            if confirmation == "yes":
                doc.add_paragraph('User inserted {} to confirm the order.'.format(confirmation))
                doc.save('regular_user_session.docx')
                print("Order confirmed!")
            elif confirmation == "no":
                doc.add_paragraph('User inserted {} to cancel order.'.format(confirmation))
                doc.save('regular_user_session.docx')
                make_order()
            else:
                print("Invalid input")
                make_order()
    else:
        print("Invalid input")
        make_order()
        
        


def make_payments(username):
    # This function will subtract money from the wallet of the customer
    # It will be the payment part after ordering
    global itemtype
    global user_choice_item
    global amount
    global doc
    global item_count_delivered
    if user_choice_item in food_class.food_item.keys():
        total_price = food_class.food_item[user_choice_item].calculate_total_price(user_choice_item, amount, itemtype)
        print("The total price is RWF" + str(total_price))
        if total_price <= regular_user[username].wallet:
            regular_user[username].subtract_money(total_price)
            user_location = regular_user[username].location
            item_class.delivery(user_location)
            doc.add_paragraph('User wallet has been subtracted by RWF {}.'.format(total_price))
            item_count_delivered += 1
            food_class.food_item[user_choice_item].item_count_delivered = item_count_delivered
            doc.save('regular_user_session.docx')
        else:
            print("Insufficient fund, cannot perform transaction!")
            doc.add_paragraph('User wallet is not sufficient to perform transaction.')
            doc.save('regular_user_session.docx')
            confirmation = input("Do you wish to re-order again? [yes/no]: ").lower()
            if confirmation == 'yes':
                doc.add_paragraph('User inserted {} to re-order again.'.format(confirmation))
                doc.save('regular_user_session.docx')
                make_order()
            else:
                doc.add_paragraph('User inserted {} to exit the program.'.format(confirmation))
                doc.save('regular_user_session.docx')
                exit()
    elif user_choice_item in drink_class.drink_item.keys():
        total_price = drink_class.drink_item[user_choice_item].calculate_total_price(user_choice_item, amount, itemtype)
        print("The total price is RWF " + str(total_price))
        if total_price <= regular_user[username].wallet:
            regular_user[username].subtract_money(total_price)
            user_location = regular_user[username].location
            item_class.delivery(user_location)
            doc.add_paragraph('User wallet has been subtracted by RWF {}.'.format(total_price))
            doc.save('regular_user_session.docx')
        else:
            print("Insufficient fund, cannot perform transaction!")
            doc.add_paragraph('User wallet is not sufficient to perform transaction.')
            doc.save('regular_user_session.docx')
            confirmation = input("Do you wish to re-order again? [yes/no]: ").lower()
            if confirmation == 'yes':
                doc.add_paragraph('User inserted {} to re-order again.'.format(confirmation))
                doc.save('regular_user_session.docx')
                make_order()
            else:
                doc.add_paragraph('User inserted {} to exit the program.'.format(confirmation))
                doc.save('regular_user_session.docx')
                exit()
    else:
        print("Error, Try again!")
        doc.add_paragraph('User get error in payment.')
        doc.save('regular_user_session.docx')
        make_order()
    


def user_review():
    # After ordering and getting delivered, it will ask the user feedback and rating
    global doc
    if user_choice_item in food_class.food_item.keys():
        user_rating_value = int(input("Please, rate the service [1-5]:  \n"))
        try:
            user_rating_value = int(user_rating_value)
        except ValueError:
            print('Only integer from 1 to 5 allowed.')
            print('Please, try again!')
            user_review()
        doc.add_paragraph('User inserted {} as rating for the service.'.format(user_rating_value))
        food_class.food_item[user_choice_item].user_rating(user_rating_value)
        user_feedback = input("Any feedback:  \n")
        doc.add_paragraph('User inserted {} as feedback.'.format(user_feedback))
        doc.save('regular_user_session.docx')
        food_class.food_item[user_choice_item].review = user_feedback
        print("Thanks for your feedback!")
        main_application.menu()
    elif user_choice_item in drink_class.drink_item.keys():
        user_rating_value = int(input("Please, rate the service [1-5]:  \n"))
        try:
            user_rating_value = int(user_rating_value)
        except ValueError:
            print('Only integer from 1 tp 5 allowed.')
            print('Please, try again!')
            user_review()
        doc.add_paragraph('User inserted {} as rating for the service.'.format(user_rating_value))
        drink_class.drink_item[user_choice_item].user_rating(user_rating_value)
        user_feedback = input("Any feedback:  \n")
        doc.add_paragraph('User inserted {} as feedback.'.format(user_feedback))
        doc.save('regular_user_session.docx')
        drink_class.drink_item[user_choice_item].review = user_feedback
        print("Thanks for your feedback!")
        main_application.menu()
    else:
        print("Error, Try again!")
        user_review()
    
